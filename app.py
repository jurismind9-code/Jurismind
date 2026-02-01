from flask import Flask, render_template, request, jsonify, session, redirect, url_for
from flask_cors import CORS
import numpy as np
from io import BytesIO
import os
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import secrets
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import traceback
import json
import random
import re
from datetime import datetime, timedelta
from functools import wraps
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# Load environment variables from .env file
load_dotenv()

# Lazy load heavy libraries
fitz = None
faiss = None

def get_fitz():
    global fitz
    if fitz is None:
        import fitz as _fitz
        fitz = _fitz
    return fitz

def get_faiss():
    global faiss
    if faiss is None:
        import faiss as _faiss
        faiss = _faiss
    return faiss

# Use google.generativeai package (more stable)
# Force using the older API which is better documented
import google.generativeai as genai
USE_NEW_GENAI = False
print("Using google.generativeai package")

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})  # Enable CORS for all routes
app.secret_key = 'jurismind-ai-secret-key-2025'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0

# Global storage (in production, use database)
vector_stores = {}
embedding_model = None
uploaded_files_info = {}
SentenceTransformer = None

# User database file path
USERS_DB_FILE = os.path.join(os.path.dirname(__file__), 'users_db.json')

# OTP storage (in production, use Redis or database)
otp_storage = {}

# Subscription Plans Configuration
SUBSCRIPTION_PLANS = {
    'free': {
        'name': 'Free',
        'price_inr': 0,
        'price_usd': 0,
        'documents_per_month': 10,
        'max_pages': 40,
        'features': {
            'case_brief': True,
            'document_analysis': False,
            'semantic_qa': False,
            'hindi_support': False,
            'url_support': False,
            'docx_support': False,
            'export_pdf': False,
            'export_docx': False,
            'history': False,
            'priority_processing': False,
            'voice_input': False
        },
        'allowed_formats': ['pdf'],
        'description': 'Perfect for getting started'
    },
    'pro': {
        'name': 'Pro',
        'price_inr': 149,
        'price_usd': 12,
        'documents_per_month': 80,
        'max_pages': 500,
        'features': {
            'case_brief': True,
            'document_analysis': True,
            'semantic_qa': True,
            'hindi_support': True,
            'url_support': True,
            'docx_support': True,
            'export_pdf': True,
            'export_docx': True,
            'history': True,
            'priority_processing': True,
            'voice_input': True
        },
        'allowed_formats': ['pdf', 'docx', 'txt'],
        'description': 'Best for professionals'
    }
}

# UPI Payment Configuration
UPI_ID = "iitianlakshya17-1@oksbi"

# Admin Configuration
ADMIN_EMAIL = "iitbhulakshya1@gmail.com"
ADMIN_PASSWORD = "Vanshika@1"

# Payment requests storage
PAYMENT_REQUESTS_FILE = os.path.join(os.path.dirname(__file__), 'payment_requests.json')
PAYMENT_SCREENSHOTS_DIR = os.path.join(os.path.dirname(__file__), 'payment_screenshots')

# Create screenshots directory if it doesn't exist
if not os.path.exists(PAYMENT_SCREENSHOTS_DIR):
    os.makedirs(PAYMENT_SCREENSHOTS_DIR)

def load_payment_requests():
    """Load payment requests from JSON file"""
    if os.path.exists(PAYMENT_REQUESTS_FILE):
        try:
            with open(PAYMENT_REQUESTS_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_payment_requests(requests):
    """Save payment requests to JSON file"""
    with open(PAYMENT_REQUESTS_FILE, 'w') as f:
        json.dump(requests, f, indent=2)

# Usage tracking file
USAGE_DB_FILE = os.path.join(os.path.dirname(__file__), 'usage_db.json')

# Document history file
HISTORY_DB_FILE = os.path.join(os.path.dirname(__file__), 'history_db.json')

# Achievements database file
ACHIEVEMENTS_DB_FILE = os.path.join(os.path.dirname(__file__), 'achievements_db.json')

# Achievement Definitions
ACHIEVEMENTS = {
    'first_upload': {
        'id': 'first_upload',
        'name': 'First Steps',
        'description': 'Upload your first document',
        'icon': 'fa-file-upload',
        'color': '#10b981',
        'requirement': 1,
        'field': 'documents_uploaded'
    },
    'docs_5': {
        'id': 'docs_5',
        'name': 'Document Explorer',
        'description': 'Upload 5 documents',
        'icon': 'fa-folder-open',
        'color': '#3b82f6',
        'requirement': 5,
        'field': 'documents_uploaded'
    },
    'docs_25': {
        'id': 'docs_25',
        'name': 'Legal Library',
        'description': 'Upload 25 documents',
        'icon': 'fa-archive',
        'color': '#8b5cf6',
        'requirement': 25,
        'field': 'documents_uploaded'
    },
    'docs_100': {
        'id': 'docs_100',
        'name': 'Document Master',
        'description': 'Upload 100 documents',
        'icon': 'fa-crown',
        'color': '#f59e0b',
        'requirement': 100,
        'field': 'documents_uploaded'
    },
    'first_brief': {
        'id': 'first_brief',
        'name': 'Brief Beginner',
        'description': 'Generate your first case brief',
        'icon': 'fa-file-alt',
        'color': '#10b981',
        'requirement': 1,
        'field': 'briefs_generated'
    },
    'briefs_10': {
        'id': 'briefs_10',
        'name': 'Brief Writer',
        'description': 'Generate 10 case briefs',
        'icon': 'fa-pen-fancy',
        'color': '#3b82f6',
        'requirement': 10,
        'field': 'briefs_generated'
    },
    'briefs_50': {
        'id': 'briefs_50',
        'name': 'Brief Expert',
        'description': 'Generate 50 case briefs',
        'icon': 'fa-graduation-cap',
        'color': '#8b5cf6',
        'requirement': 50,
        'field': 'briefs_generated'
    },
    'first_qa': {
        'id': 'first_qa',
        'name': 'Curious Mind',
        'description': 'Ask your first question',
        'icon': 'fa-question-circle',
        'color': '#10b981',
        'requirement': 1,
        'field': 'qa_queries'
    },
    'qa_25': {
        'id': 'qa_25',
        'name': 'Knowledge Seeker',
        'description': 'Ask 25 questions',
        'icon': 'fa-search',
        'color': '#3b82f6',
        'requirement': 25,
        'field': 'qa_queries'
    },
    'qa_100': {
        'id': 'qa_100',
        'name': 'Legal Scholar',
        'description': 'Ask 100 questions',
        'icon': 'fa-user-graduate',
        'color': '#8b5cf6',
        'requirement': 100,
        'field': 'qa_queries'
    },
    'qa_500': {
        'id': 'qa_500',
        'name': 'Research Guru',
        'description': 'Ask 500 questions',
        'icon': 'fa-star',
        'color': '#f59e0b',
        'requirement': 500,
        'field': 'qa_queries'
    },
    'pages_100': {
        'id': 'pages_100',
        'name': 'Page Turner',
        'description': 'Process 100 pages',
        'icon': 'fa-book-open',
        'color': '#3b82f6',
        'requirement': 100,
        'field': 'pages_processed'
    },
    'pages_1000': {
        'id': 'pages_1000',
        'name': 'Bookworm',
        'description': 'Process 1000 pages',
        'icon': 'fa-books',
        'color': '#8b5cf6',
        'requirement': 1000,
        'field': 'pages_processed'
    }
}

# Daily Legal Tips
DAILY_TIPS = [
    {
        'title': 'Know Your Rights',
        'content': 'Under Article 21 of the Indian Constitution, every person has the right to life and personal liberty. This is one of the most fundamental rights.',
        'category': 'Constitutional Law',
        'icon': 'fa-balance-scale'
    },
    {
        'title': 'FIR Filing',
        'content': 'You can file a Zero FIR at any police station in India, regardless of jurisdiction. The police are obligated to register it.',
        'category': 'Criminal Law',
        'icon': 'fa-gavel'
    },
    {
        'title': 'Consumer Protection',
        'content': 'Under the Consumer Protection Act 2019, you can file complaints online through the e-Daakhil portal for amounts up to Rs. 1 crore.',
        'category': 'Consumer Law',
        'icon': 'fa-shield-alt'
    },
    {
        'title': 'RTI Power',
        'content': 'Under the RTI Act, public authorities must respond to your information request within 30 days. For life and liberty matters, it\'s 48 hours.',
        'category': 'RTI Act',
        'icon': 'fa-info-circle'
    },
    {
        'title': 'Anticipatory Bail',
        'content': 'Section 438 of CrPC allows you to apply for anticipatory bail if you apprehend arrest in a non-bailable offense.',
        'category': 'Criminal Procedure',
        'icon': 'fa-unlock'
    },
    {
        'title': 'Cyber Crime Reporting',
        'content': 'Report cyber crimes at cybercrime.gov.in. Financial frauds should be reported within the "golden hour" for better recovery chances.',
        'category': 'Cyber Law',
        'icon': 'fa-laptop-code'
    },
    {
        'title': 'Tenant Rights',
        'content': 'A landlord cannot forcefully evict a tenant. Proper legal notice and court procedure must be followed as per the Rent Control Act.',
        'category': 'Property Law',
        'icon': 'fa-home'
    },
    {
        'title': 'Witness Protection',
        'content': 'India has a Witness Protection Scheme 2018 that provides security to witnesses in criminal cases if they face threats.',
        'category': 'Criminal Law',
        'icon': 'fa-user-shield'
    },
    {
        'title': 'Legal Aid',
        'content': 'Free legal aid is a constitutional right under Article 39A. Contact NALSA or your State Legal Services Authority.',
        'category': 'Legal Aid',
        'icon': 'fa-hands-helping'
    },
    {
        'title': 'Cheque Bounce',
        'content': 'Under Section 138 of NI Act, cheque bounce is a criminal offense. Send a legal notice within 30 days of receiving the memo.',
        'category': 'Banking Law',
        'icon': 'fa-money-check-alt'
    },
    {
        'title': 'Marriage Registration',
        'content': 'Marriage registration is compulsory in many states. It serves as conclusive proof of marriage under the law.',
        'category': 'Family Law',
        'icon': 'fa-ring'
    },
    {
        'title': 'Workplace Harassment',
        'content': 'Every organization with 10+ employees must have an Internal Complaints Committee under the POSH Act 2013.',
        'category': 'Employment Law',
        'icon': 'fa-briefcase'
    },
    {
        'title': 'Bail is the Rule',
        'content': 'The Supreme Court has repeatedly held that "Bail is the rule, jail is the exception" for non-heinous offenses.',
        'category': 'Criminal Law',
        'icon': 'fa-balance-scale-right'
    },
    {
        'title': 'E-Courts Services',
        'content': 'Check your case status online at ecourts.gov.in. You can also pay court fees and get certified copies digitally.',
        'category': 'Court Procedure',
        'icon': 'fa-desktop'
    },
    {
        'title': 'Mediation First',
        'content': 'Under Section 89 of CPC, courts can refer disputes to mediation. It\'s faster, cheaper, and maintains relationships.',
        'category': 'ADR',
        'icon': 'fa-handshake'
    }
]

def load_achievements_db():
    """Load achievements data from JSON file"""
    if os.path.exists(ACHIEVEMENTS_DB_FILE):
        try:
            with open(ACHIEVEMENTS_DB_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_achievements_db(achievements):
    """Save achievements data to JSON file"""
    with open(ACHIEVEMENTS_DB_FILE, 'w') as f:
        json.dump(achievements, f, indent=2)

def get_user_achievements(user_id):
    """Get user's unlocked achievements"""
    achievements = load_achievements_db()
    return achievements.get(user_id, {'unlocked': [], 'notified': []})

def check_and_unlock_achievements(user_id, total_stats):
    """Check if user has unlocked new achievements based on their stats"""
    user_achievements = get_user_achievements(user_id)
    unlocked = user_achievements.get('unlocked', [])
    notified = user_achievements.get('notified', [])
    new_achievements = []

    for ach_id, achievement in ACHIEVEMENTS.items():
        if ach_id not in unlocked:
            field = achievement['field']
            requirement = achievement['requirement']
            current_value = total_stats.get(field, 0)

            if current_value >= requirement:
                unlocked.append(ach_id)
                new_achievements.append(achievement)

    # Save updated achievements
    if new_achievements:
        achievements_db = load_achievements_db()
        achievements_db[user_id] = {'unlocked': unlocked, 'notified': notified}
        save_achievements_db(achievements_db)

    return new_achievements

def get_total_user_stats(user_id):
    """Get total stats for a user across all time"""
    usage = load_usage_db()
    total_stats = {
        'documents_uploaded': 0,
        'pages_processed': 0,
        'briefs_generated': 0,
        'analyses_done': 0,
        'qa_queries': 0
    }

    if user_id in usage:
        for month, stats in usage[user_id].items():
            for key in total_stats:
                total_stats[key] += stats.get(key, 0)

    return total_stats

def get_daily_tip():
    """Get a daily tip based on current date"""
    day_of_year = datetime.now().timetuple().tm_yday
    tip_index = day_of_year % len(DAILY_TIPS)
    return DAILY_TIPS[tip_index]

def load_usage_db():
    """Load usage data from JSON file"""
    if os.path.exists(USAGE_DB_FILE):
        try:
            with open(USAGE_DB_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_usage_db(usage):
    """Save usage data to JSON file"""
    with open(USAGE_DB_FILE, 'w') as f:
        json.dump(usage, f, indent=2)

def load_history_db():
    """Load document history from JSON file"""
    if os.path.exists(HISTORY_DB_FILE):
        try:
            with open(HISTORY_DB_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_history_db(history):
    """Save document history to JSON file"""
    with open(HISTORY_DB_FILE, 'w') as f:
        json.dump(history, f, indent=2)

def get_user_plan(user_id):
    """Get user's current subscription plan"""
    users = load_users_db()
    if user_id in users:
        return users[user_id].get('plan', 'free')
    return 'free'

def get_user_usage(user_id):
    """Get user's usage for current month"""
    usage = load_usage_db()
    current_month = datetime.now().strftime('%Y-%m')

    if user_id not in usage:
        usage[user_id] = {}

    if current_month not in usage[user_id]:
        usage[user_id][current_month] = {
            'documents_uploaded': 0,
            'pages_processed': 0,
            'briefs_generated': 0,
            'analyses_done': 0,
            'qa_queries': 0
        }
        save_usage_db(usage)

    return usage[user_id][current_month]

def increment_usage(user_id, field, amount=1):
    """Increment a usage field for the user"""
    usage = load_usage_db()
    current_month = datetime.now().strftime('%Y-%m')

    if user_id not in usage:
        usage[user_id] = {}

    if current_month not in usage[user_id]:
        usage[user_id][current_month] = {
            'documents_uploaded': 0,
            'pages_processed': 0,
            'briefs_generated': 0,
            'analyses_done': 0,
            'qa_queries': 0
        }

    usage[user_id][current_month][field] = usage[user_id][current_month].get(field, 0) + amount
    save_usage_db(usage)
    return usage[user_id][current_month]

def check_usage_limit(user_id, check_type='documents'):
    """Check if user has exceeded their plan limits"""
    plan_name = get_user_plan(user_id)
    plan = SUBSCRIPTION_PLANS[plan_name]
    usage = get_user_usage(user_id)

    if check_type == 'documents':
        limit = plan['documents_per_month']
        current = usage.get('documents_uploaded', 0)
    elif check_type == 'pages':
        limit = plan['max_pages']
        current = usage.get('pages_processed', 0)
    else:
        return True, 0, -1

    if limit == -1:  # Unlimited
        return True, current, -1

    return current < limit, current, limit

def check_feature_access(user_id, feature):
    """Check if user has access to a specific feature"""
    plan_name = get_user_plan(user_id)
    plan = SUBSCRIPTION_PLANS[plan_name]
    return plan['features'].get(feature, False)

def add_to_history(user_id, doc_info):
    """Add document to user's history"""
    history = load_history_db()

    if user_id not in history:
        history[user_id] = []

    doc_entry = {
        'id': secrets.token_hex(8),
        'filename': doc_info.get('filename'),
        'upload_date': datetime.now().isoformat(),
        'pages': doc_info.get('pages', 0),
        'language': doc_info.get('language', 'en'),
        'brief_generated': doc_info.get('brief_generated', False),
        'analysis_done': doc_info.get('analysis_done', False),
        'case_brief': None,  # Will store the generated case brief
        'chat_history': []   # Will store Q&A conversations
    }

    history[user_id].insert(0, doc_entry)  # Add to beginning

    # Keep only last 100 entries
    history[user_id] = history[user_id][:100]

    save_history_db(history)
    return doc_entry

def update_history_brief(user_id, doc_id, brief_content):
    """Update history entry with generated case brief"""
    history = load_history_db()

    if user_id not in history:
        return False

    for entry in history[user_id]:
        if entry.get('id') == doc_id:
            entry['case_brief'] = brief_content
            entry['brief_generated'] = True
            entry['brief_generated_at'] = datetime.now().isoformat()
            save_history_db(history)
            return True

    return False

def add_chat_to_history(user_id, doc_id, question, answer):
    """Add a Q&A exchange to document's history"""
    history = load_history_db()

    if user_id not in history:
        return False

    for entry in history[user_id]:
        if entry.get('id') == doc_id:
            if 'chat_history' not in entry:
                entry['chat_history'] = []

            entry['chat_history'].append({
                'question': question,
                'answer': answer,
                'timestamp': datetime.now().isoformat()
            })

            # Keep only last 50 Q&A per document
            entry['chat_history'] = entry['chat_history'][-50:]
            save_history_db(history)
            return True

    return False

def get_current_doc_id(user_id):
    """Get the ID of the most recently uploaded document"""
    history = load_history_db()
    if user_id in history and history[user_id]:
        return history[user_id][0].get('id')
    return None

def get_user_history(user_id, limit=20):
    """Get user's document history"""
    history = load_history_db()
    return history.get(user_id, [])[:limit]

def load_users_db():
    """Load users from JSON file"""
    if os.path.exists(USERS_DB_FILE):
        try:
            with open(USERS_DB_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_users_db(users):
    """Save users to JSON file"""
    with open(USERS_DB_FILE, 'w') as f:
        json.dump(users, f, indent=2)

def generate_otp():
    """Generate a 6-digit OTP"""
    return str(random.randint(100000, 999999))

def send_otp_sms(phone, otp):
    """
    Send OTP via SMS. In production, integrate with SMS gateway like:
    - Twilio
    - MSG91
    - AWS SNS
    For demo purposes, we'll just print/log the OTP
    """
    print(f"[OTP] Sending OTP {otp} to phone {phone}")
    # In production, uncomment and configure your SMS provider:
    # Example with Twilio:
    # from twilio.rest import Client
    # client = Client(account_sid, auth_token)
    # message = client.messages.create(
    #     body=f"Your JurisMind AI OTP is: {otp}. Valid for 5 minutes.",
    #     from_='+1234567890',
    #     to=phone
    # )
    return True

def validate_phone(phone):
    """Validate phone number format"""
    # Remove spaces and dashes
    phone = re.sub(r'[\s\-]', '', phone)
    # Check if it's a valid phone number (10+ digits)
    if re.match(r'^\+?\d{10,15}$', phone):
        return phone
    return None

def validate_email(email):
    """Validate email format"""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def login_required(f):
    """Decorator to require authentication"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': 'Authentication required', 'redirect': '/login'}), 401
        return f(*args, **kwargs)
    return decorated_function

def get_embedding_model():
    global embedding_model, SentenceTransformer
    if embedding_model is None:
        print("Loading embedding model...")
        try:
            if SentenceTransformer is None:
                from sentence_transformers import SentenceTransformer as ST
                SentenceTransformer = ST
            embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            print("Embedding model loaded!")
        except Exception as e:
            print(f"Error loading embedding model: {e}")
            traceback.print_exc()
            raise
    return embedding_model

# OCR Engine configuration - Multilingual support (English + Hindi)
easyocr_reader = None

def get_easyocr_reader():
    """Lazy load EasyOCR reader with Hindi and English support"""
    global easyocr_reader
    if easyocr_reader is None:
        try:
            # Fix Windows encoding issue
            import sys
            import io
            if sys.platform == 'win32':
                sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
                sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

            import easyocr
            print("Loading EasyOCR model with Hindi + English support...")
            # Load both English and Hindi language models
            easyocr_reader = easyocr.Reader(['en', 'hi'], gpu=False, verbose=False)
            print("EasyOCR loaded successfully with Hindi + English support!")
        except Exception as e:
            print(f"Failed to load EasyOCR: {e}")
            return None
    return easyocr_reader

def detect_language(text):
    """Detect if text is primarily Hindi or English"""
    if not text:
        return 'en'

    # Hindi Unicode range: \u0900-\u097F (Devanagari)
    hindi_chars = len(re.findall(r'[\u0900-\u097F]', text))
    english_chars = len(re.findall(r'[a-zA-Z]', text))

    total = hindi_chars + english_chars
    if total == 0:
        return 'en'

    hindi_ratio = hindi_chars / total

    if hindi_ratio > 0.3:  # More than 30% Hindi characters
        return 'hi'
    return 'en'

def extract_text_from_image_ocr(pil_image):
    """Extract text from a PIL Image using EasyOCR (Hindi + English)"""
    try:
        import numpy as np
        reader = get_easyocr_reader()
        if reader is None:
            return "", "en"

        # Convert PIL Image to numpy array
        img_array = np.array(pil_image)

        # Perform OCR with both languages
        results = reader.readtext(img_array)

        # Extract text from results
        text_parts = [result[1] for result in results]
        text = ' '.join(text_parts)

        # Detect language
        detected_lang = detect_language(text)

        return text.strip(), detected_lang
    except Exception as e:
        print(f"OCR error: {e}")
        return ""

def extract_text_from_pdf(file_bytes):
    """Extract text from PDF with OCR support for scanned/image-based PDFs (Hindi + English)"""
    text = ""
    ocr_used = False
    detected_language = "en"

    try:
        pdf_lib = get_fitz()
        print(f"Opening PDF ({len(file_bytes)} bytes)...")

        with pdf_lib.open(stream=file_bytes, filetype="pdf") as doc:
            total_pages = len(doc)
            print(f"PDF has {total_pages} pages")

            for page_num, page in enumerate(doc):
                try:
                    # First try to get text directly
                    page_text = page.get_text().strip()

                    # If no text found, try OCR
                    if not page_text or len(page_text) < 50:
                        print(f"Page {page_num + 1}: Little/no text ({len(page_text)} chars), trying OCR...")
                        ocr_used = True

                        try:
                            # Render page to image at 150 DPI (balance quality/speed)
                            mat = pdf_lib.Matrix(150/72, 150/72)
                            pix = page.get_pixmap(matrix=mat)

                            # Convert to PIL Image
                            from PIL import Image
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                            # Perform OCR (returns text and detected language)
                            ocr_result = extract_text_from_image_ocr(img)
                            if isinstance(ocr_result, tuple):
                                ocr_text, lang = ocr_result
                                if lang == 'hi':
                                    detected_language = 'hi'
                            else:
                                ocr_text = ocr_result

                            if ocr_text:
                                page_text = ocr_text
                                print(f"Page {page_num + 1}: OCR extracted {len(ocr_text)} chars")

                        except Exception as ocr_err:
                            print(f"OCR failed on page {page_num + 1}: {ocr_err}")

                    # Also try to extract text from embedded images in the page
                    if not page_text or len(page_text) < 50:
                        try:
                            image_list = page.get_images()
                            for img_index, img_info in enumerate(image_list[:3]):  # Limit to first 3 images
                                try:
                                    xref = img_info[0]
                                    base_image = doc.extract_image(xref)
                                    image_bytes = base_image["image"]

                                    from PIL import Image
                                    pil_img = Image.open(BytesIO(image_bytes))

                                    # Only OCR reasonably sized images
                                    if pil_img.width > 100 and pil_img.height > 100:
                                        ocr_result = extract_text_from_image_ocr(pil_img)
                                        if isinstance(ocr_result, tuple):
                                            img_text, lang = ocr_result
                                            if lang == 'hi':
                                                detected_language = 'hi'
                                        else:
                                            img_text = ocr_result
                                        if img_text:
                                            page_text += " " + img_text
                                except:
                                    continue
                        except:
                            pass

                    text += page_text + "\n"

                    # Progress logging
                    if (page_num + 1) % 5 == 0 or page_num == 0:
                        print(f"Processed {page_num + 1}/{total_pages} pages...")

                except Exception as page_error:
                    print(f"Error on page {page_num + 1}: {page_error}")
                    continue

            if ocr_used:
                print(f"OCR was used for some pages")

            # Also detect language from extracted text
            if detected_language == 'en':
                detected_language = detect_language(text)

            print(f"Extracted {len(text)} characters from {total_pages} pages")
            print(f"Detected language: {'Hindi' if detected_language == 'hi' else 'English'}")

    except Exception as e:
        print(f"Error extracting PDF: {e}")
        traceback.print_exc()

    return text, detected_language

def extract_text_from_docx(file_bytes):
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        detected_lang = detect_language(text)
        return text, detected_lang
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        traceback.print_exc()
        return "", "en"

def extract_text_from_url(url):
    """Extract text content from a URL"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()

        # Get text
        text = soup.get_text(separator='\n')

        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        return text
    except Exception as e:
        print(f"Error extracting URL content: {e}")
        traceback.print_exc()
        return ""

def chunk_text(text, chunk_size=500, overlap=100):
    """Split text into overlapping chunks - optimized for large documents"""
    if not text or not text.strip():
        return []

    # Clean and normalize text
    text = ' '.join(text.split())  # Normalize whitespace

    words = text.split()
    if len(words) == 0:
        return []

    print(f"Chunking {len(words)} words into chunks of ~{chunk_size} words...")

    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip() and len(chunk) > 50:  # Minimum chunk size
            chunks.append(chunk)

    # Ensure at least one chunk
    if not chunks and text.strip():
        chunks = [text[:5000]]  # Take first 5000 chars if chunking fails

    print(f"Created {len(chunks)} chunks")
    return chunks

def get_available_models(api_key):
    """Get list of available models that support generateContent"""
    try:
        genai.configure(api_key=api_key)
        available = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                available.append(m.name)
        return available
    except Exception as e:
        print(f"Error listing models: {e}")
        return []

def gemini_response(prompt, api_key):
    """Get response from Gemini API"""
    try:
        # Configure the API
        genai.configure(api_key=api_key)

        # First, try to get available models
        available_models = get_available_models(api_key)
        if available_models:
            print(f"Available models: {available_models[:5]}...")  # Print first 5

        # Model names to try - with models/ prefix as required by API
        models_to_try = [
            'models/gemini-2.0-flash-exp',
            'models/gemini-1.5-flash-latest',
            'models/gemini-1.5-flash',
            'models/gemini-1.5-pro-latest',
            'models/gemini-1.5-pro',
            'models/gemini-1.0-pro',
            'models/gemini-pro',
        ]

        # Also try without prefix (some SDK versions accept both)
        models_to_try_alt = [
            'gemini-2.0-flash-exp',
            'gemini-1.5-flash-latest',
            'gemini-1.5-flash',
            'gemini-1.5-pro-latest',
            'gemini-1.5-pro',
            'gemini-1.0-pro',
            'gemini-pro',
        ]

        # Try models with prefix first
        for model_name in models_to_try:
            try:
                print(f"Trying model: {model_name}")
                model = genai.GenerativeModel(model_name)
                response = model.generate_content(prompt)
                if response and response.text:
                    print(f"Success with model: {model_name}")
                    return response.text
            except Exception as e:
                error_str = str(e)
                print(f"Error with {model_name}: {error_str[:100]}")
                continue

        # Try without prefix
        for model_name in models_to_try_alt:
            try:
                print(f"Trying model (alt): {model_name}")
                model = genai.GenerativeModel(model_name)
                response = model.generate_content(prompt)
                if response and response.text:
                    print(f"Success with model: {model_name}")
                    return response.text
            except Exception as e:
                error_str = str(e)
                print(f"Error with {model_name}: {error_str[:100]}")
                continue

        # Last resort: try any available model from the list
        if available_models:
            for model_name in available_models:
                if 'gemini' in model_name.lower():
                    try:
                        print(f"Trying available model: {model_name}")
                        model = genai.GenerativeModel(model_name)
                        response = model.generate_content(prompt)
                        if response and response.text:
                            print(f"Success with model: {model_name}")
                            return response.text
                    except Exception as e:
                        continue

        print("All models failed!")
        return None
    except Exception as e:
        print(f"Gemini API error: {e}")
        traceback.print_exc()
        return None

def get_or_create_session_id():
    """Get existing session ID or create new one"""
    if 'session_id' not in session:
        session['session_id'] = secrets.token_hex(16)
    return session['session_id']

def add_to_vector_store(session_id, text, source_name, language='en'):
    """Add text to the vector store for a session - optimized for large docs"""
    if not text or not text.strip():
        print(f"No text to add for {source_name}")
        return 0

    print(f"Adding {source_name} to vector store ({len(text)} chars, language: {language})...")

    chunks = chunk_text(text)
    if not chunks:
        print(f"No chunks created for {source_name}")
        return 0

    print(f"Encoding {len(chunks)} chunks...")
    model = get_embedding_model()

    # Process in batches for large documents
    batch_size = 32
    all_embeddings = []

    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i + batch_size]
        batch_embeddings = model.encode(batch, show_progress_bar=False)
        all_embeddings.append(batch_embeddings)
        if (i + batch_size) % 100 == 0:
            print(f"Encoded {min(i + batch_size, len(chunks))}/{len(chunks)} chunks...")

    embeddings = np.vstack(all_embeddings).astype('float32')
    print(f"Created embeddings: shape {embeddings.shape}")

    faiss_lib = get_faiss()

    if session_id not in vector_stores:
        # Create new store
        index = faiss_lib.IndexFlatL2(embeddings.shape[1])
        vector_stores[session_id] = {
            'index': index,
            'chunks': [],
            'texts': {},
            'sources': [],
            'languages': {}
        }

    store = vector_stores[session_id]
    store['index'].add(embeddings)
    store['chunks'].extend(chunks)
    store['texts'][source_name] = text
    store['sources'].append(source_name)

    # Store language info for the document
    if 'languages' not in store:
        store['languages'] = {}
    store['languages'][source_name] = language

    print(f"Successfully added {len(chunks)} chunks for {source_name} (language: {language})")
    return len(chunks)

@app.route('/')
def index():
    """Main page - requires authentication"""
    if 'user_id' not in session:
        return redirect('/login')
    # Auto-set API key from .env if not already set
    if 'api_key' not in session:
        env_api_key = os.getenv('GEMINI_API_KEY')
        if env_api_key:
            session['api_key'] = env_api_key
    return render_template('index.html')

@app.route('/login')
def login_page():
    """Login page"""
    if 'user_id' in session:
        return redirect('/')
    return render_template('login.html')

@app.route('/signup')
def signup_page():
    """Signup page"""
    if 'user_id' in session:
        return redirect('/')
    return render_template('signup.html')

@app.route('/api/auth/signup', methods=['POST'])
def signup():
    """Handle user signup - Step 1: Collect info and send OTP"""
    try:
        data = request.json
        name = data.get('name', '').strip()
        email = data.get('email', '').strip().lower()
        phone = data.get('phone', '').strip()
        password = data.get('password', '').strip()

        # Validation
        if not name or len(name) < 2:
            return jsonify({'success': False, 'message': 'Please enter a valid name'})

        if not email or not validate_email(email):
            return jsonify({'success': False, 'message': 'Please enter a valid email address'})

        phone = validate_phone(phone)
        if not phone:
            return jsonify({'success': False, 'message': 'Please enter a valid phone number (10+ digits)'})

        if not password or len(password) < 6:
            return jsonify({'success': False, 'message': 'Password must be at least 6 characters'})

        # Check if user already exists
        users = load_users_db()
        for user_id, user in users.items():
            if user['email'] == email:
                return jsonify({'success': False, 'message': 'Email already registered. Please login.'})
            if user['phone'] == phone:
                return jsonify({'success': False, 'message': 'Phone number already registered. Please login.'})

        # Generate and store OTP
        otp = generate_otp()
        otp_storage[phone] = {
            'otp': otp,
            'name': name,
            'email': email,
            'phone': phone,
            'password': generate_password_hash(password, method='pbkdf2:sha256'),
            'expires': (datetime.now() + timedelta(minutes=5)).isoformat()
        }

        # Send OTP (in production, this would send actual SMS)
        send_otp_sms(phone, otp)

        # For demo/testing, also return OTP in response (remove in production!)
        return jsonify({
            'success': True,
            'message': f'OTP sent to {phone}. Please verify.',
            'demo_otp': otp  # Remove this line in production!
        })

    except Exception as e:
        print(f"Signup error: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/auth/verify-otp', methods=['POST'])
def verify_otp():
    """Verify OTP and complete signup"""
    try:
        data = request.json
        phone = validate_phone(data.get('phone', '').strip())
        otp = data.get('otp', '').strip()

        if not phone or phone not in otp_storage:
            return jsonify({'success': False, 'message': 'Invalid phone number or OTP expired'})

        stored = otp_storage[phone]

        # Check expiry
        if datetime.now() > datetime.fromisoformat(stored['expires']):
            del otp_storage[phone]
            return jsonify({'success': False, 'message': 'OTP expired. Please signup again.'})

        # Verify OTP
        if stored['otp'] != otp:
            return jsonify({'success': False, 'message': 'Invalid OTP. Please try again.'})

        # Create user
        users = load_users_db()
        user_id = secrets.token_hex(16)

        users[user_id] = {
            'id': user_id,
            'name': stored['name'],
            'email': stored['email'],
            'phone': stored['phone'],
            'password': stored['password'],
            'created_at': datetime.now().isoformat(),
            'verified': True
        }

        save_users_db(users)

        # Clean up OTP
        del otp_storage[phone]

        # Auto login
        session['user_id'] = user_id
        session['user_name'] = stored['name']
        session['user_email'] = stored['email']

        return jsonify({
            'success': True,
            'message': 'Account created successfully!',
            'user': {
                'name': stored['name'],
                'email': stored['email']
            }
        })

    except Exception as e:
        print(f"OTP verification error: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/auth/resend-otp', methods=['POST'])
def resend_otp():
    """Resend OTP"""
    try:
        data = request.json
        phone = validate_phone(data.get('phone', '').strip())

        if not phone or phone not in otp_storage:
            return jsonify({'success': False, 'message': 'Please start signup again'})

        # Generate new OTP
        otp = generate_otp()
        otp_storage[phone]['otp'] = otp
        otp_storage[phone]['expires'] = (datetime.now() + timedelta(minutes=5)).isoformat()

        send_otp_sms(phone, otp)

        return jsonify({
            'success': True,
            'message': 'OTP resent successfully',
            'demo_otp': otp  # Remove in production!
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/auth/login', methods=['POST'])
def login():
    """Handle user login"""
    try:
        data = request.json
        identifier = data.get('identifier', '').strip().lower()  # Email or phone
        password = data.get('password', '').strip()

        if not identifier or not password:
            return jsonify({'success': False, 'message': 'Please enter email/phone and password'})

        users = load_users_db()

        # Find user by email or phone
        found_user = None
        for user_id, user in users.items():
            if user['email'] == identifier or user['phone'] == identifier or user['phone'] == validate_phone(identifier):
                found_user = user
                break

        if not found_user:
            return jsonify({'success': False, 'message': 'Account not found. Please signup first.'})

        # Verify password
        try:
            if not check_password_hash(found_user['password'], password):
                return jsonify({'success': False, 'message': 'Incorrect password'})
        except AttributeError as e:
            if 'scrypt' in str(e):
                return jsonify({
                    'success': False,
                    'message': 'Your account uses an outdated password format. Please use the "Forgot Password" link to reset it.'
                })
            else:
                # Re-raise other attribute errors
                raise

        # Set session
        session['user_id'] = found_user['id']
        session['user_name'] = found_user['name']
        session['user_email'] = found_user['email']

        return jsonify({
            'success': True,
            'message': f'Welcome back, {found_user["name"]}!',
            'user': {
                'name': found_user['name'],
                'email': found_user['email']
            }
        })

    except Exception as e:
        print(f"Login error: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    """Handle user logout"""
    session.clear()
    return jsonify({'success': True, 'message': 'Logged out successfully'})

@app.route('/api/auth/user', methods=['GET'])
def get_current_user():
    """Get current logged in user info"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'logged_in': False})

    return jsonify({
        'success': True,
        'logged_in': True,
        'user': {
            'name': session.get('user_name'),
            'email': session.get('user_email')
        }
    })

# Email and SMS sending functions
def send_otp_email(email, otp):
    """Send OTP via email"""
    try:
        # Email configuration (you'll need to set these in .env file)
        smtp_server = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
        smtp_port = int(os.getenv('SMTP_PORT', '587'))
        sender_email = os.getenv('SENDER_EMAIL', '')
        sender_password = os.getenv('SENDER_PASSWORD', '')
        
        if not sender_email or not sender_password:
            print("Email credentials not configured. Skipping email send.")
            return False
        
        message = MIMEMultipart()
        message['From'] = sender_email
        message['To'] = email
        message['Subject'] = 'JurisMind AI - Password Reset OTP'
        
        body = f"""
        Dear User,
        
        Your OTP for password reset is: {otp}
        
        This OTP is valid for 10 minutes only.
        
        If you didn't request this, please ignore this email.
        
        Thanks,
        JurisMind AI Team
        """
        
        message.attach(MIMEText(body, 'plain'))
        
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, sender_password)
        text = message.as_string()
        server.sendmail(sender_email, email, text)
        server.quit()
        
        print(f"OTP sent successfully to email: {email}")
        return True
        
    except Exception as e:
        print(f"Failed to send email: {str(e)}")
        return False

def send_otp_sms(phone, otp):
    """Send OTP via SMS (using a simple SMS service)"""
    try:
        # For demo purposes, we'll use a free SMS service
        # In production, you'd use services like Twilio, MessageBird, etc.
        
        # Remove any non-digits from phone number
        phone_clean = re.sub(r'\D', '', phone)
        
        # For demo, we'll just print the SMS (you can integrate with real SMS service)
        sms_message = f"Your JurisMind AI password reset OTP is: {otp}. Valid for 10 minutes."
        
        # Example with a free SMS service (you'd need to sign up)
        # This is just a placeholder - replace with actual SMS service
        sms_url = f"https://api.textlocal.in/send/?apikey=YOUR_API_KEY&numbers={phone_clean}&message={sms_message}"
        
        # For now, just print (in production, make actual API call)
        print(f"SMS would be sent to {phone}: {sms_message}")
        print(f"To enable SMS, configure SMS service API in send_otp_sms function")
        
        return True
        
    except Exception as e:
        print(f"Failed to send SMS: {str(e)}")
        return False

@app.route('/api/auth/forgot-password', methods=['POST'])
def forgot_password():
    """Send OTP for password reset"""
    try:
        data = request.json
        identifier = data.get('identifier', '').strip()

        if not identifier:
            return jsonify({'success': False, 'message': 'Please enter email or phone number'})

        users = load_users_db()

        # Find user by email or phone
        user_id = None
        user_info = None
        for uid, uinfo in users.items():
            if uinfo.get('email', '').lower() == identifier.lower() or uinfo.get('phone') == identifier:
                user_id = uid
                user_info = uinfo
                break

        if not user_id:
            return jsonify({'success': False, 'message': 'No account found with this email/phone'})

        # Generate OTP
        otp = str(random.randint(100000, 999999))
        otp_storage[identifier] = {
            'otp': otp,
            'user_id': user_id,
            'type': 'reset',
            'expires': datetime.now() + timedelta(minutes=10)
        }

        # Send OTP via email or SMS
        email_sent = False
        sms_sent = False
        
        if '@' in identifier:  # Email address
            email_sent = send_otp_email(identifier, otp)
        else:  # Phone number
            sms_sent = send_otp_sms(identifier, otp)
        
        # Always print OTP for development (as backup)
        print(f"\n{'='*50}")
        print(f"PASSWORD RESET OTP for {identifier}: {otp}")
        print(f"Email sent: {email_sent}")
        print(f"SMS sent: {sms_sent}")
        print(f"{'='*50}\n")

        return jsonify({
            'success': True,
            'message': 'OTP sent to your email/phone',
            'dev_otp': otp  # Remove in production
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/auth/verify-reset-otp', methods=['POST'])
def verify_reset_otp():
    """Verify OTP for password reset"""
    try:
        data = request.json
        identifier = data.get('identifier', '').strip()
        otp = data.get('otp', '').strip()

        if not identifier or not otp:
            return jsonify({'success': False, 'message': 'Invalid request'})

        stored = otp_storage.get(identifier)

        if not stored or stored.get('type') != 'reset':
            return jsonify({'success': False, 'message': 'Please request a new OTP'})

        if datetime.now() > stored['expires']:
            del otp_storage[identifier]
            return jsonify({'success': False, 'message': 'OTP expired. Please request a new one.'})

        if stored['otp'] != otp:
            return jsonify({'success': False, 'message': 'Invalid OTP'})

        # Mark OTP as verified for password reset
        otp_storage[identifier]['verified'] = True

        return jsonify({'success': True, 'message': 'OTP verified successfully'})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/auth/reset-password', methods=['POST'])
def reset_password():
    """Reset user password after OTP verification"""
    try:
        data = request.json
        identifier = data.get('identifier', '').strip()
        new_password = data.get('new_password', '')

        if not identifier or not new_password:
            return jsonify({'success': False, 'message': 'Invalid request'})

        if len(new_password) < 6:
            return jsonify({'success': False, 'message': 'Password must be at least 6 characters'})

        stored = otp_storage.get(identifier)

        if not stored or not stored.get('verified'):
            return jsonify({'success': False, 'message': 'Please verify OTP first'})

        # Update password
        users = load_users_db()
        user_id = stored['user_id']

        if user_id in users:
            users[user_id]['password'] = generate_password_hash(new_password, method='pbkdf2:sha256')
            save_users_db(users)

            # Clear OTP storage
            del otp_storage[identifier]

            return jsonify({'success': True, 'message': 'Password reset successfully'})
        else:
            return jsonify({'success': False, 'message': 'User not found'})

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/set-key', methods=['POST'])
def set_api_key():
    try:
        data = request.json
        api_key = data.get('api_key', '').strip()

        if not api_key:
            return jsonify({'success': False, 'message': 'No API key provided'})

        session['api_key'] = api_key

        # Quick validation test - list available models first
        genai.configure(api_key=api_key)

        # Try to list models to validate the key
        try:
            available = get_available_models(api_key)
            if available:
                print(f"API key valid! Available models: {available[:3]}...")
                # Try to actually use a model
                for model_name in available:
                    if 'gemini' in model_name.lower():
                        try:
                            model = genai.GenerativeModel(model_name)
                            response = model.generate_content("Say OK")
                            if response and response.text:
                                return jsonify({'success': True, 'message': f'API key validated with {model_name}!'})
                        except:
                            continue
                return jsonify({'success': True, 'message': f'API key valid! Found {len(available)} models.'})
        except Exception as e:
            print(f"Could not list models: {e}")

        # Fallback: try specific model names
        validation_models = [
            'models/gemini-1.5-flash',
            'models/gemini-1.5-pro',
            'models/gemini-1.0-pro',
            'gemini-1.5-flash',
            'gemini-pro'
        ]

        for model_name in validation_models:
            try:
                print(f"Validating API key with model: {model_name}")
                model = genai.GenerativeModel(model_name)
                response = model.generate_content("Say 'OK'")
                if response and response.text:
                    print(f"API key validated with model: {model_name}")
                    return jsonify({'success': True, 'message': f'API key validated with {model_name}!'})
            except Exception as e:
                print(f"Validation failed with {model_name}: {str(e)[:80]}")
                continue

        # If all models failed for validation, still save the key
        print("All validation attempts failed, but saving key anyway")
        return jsonify({'success': True, 'message': 'API key saved (validation skipped)'})
    except Exception as e:
        print(f"Error in set_api_key: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/upload', methods=['POST'])
def upload_document():
    """Handle single or multiple file uploads - optimized for large files"""
    try:
        session_id = get_or_create_session_id()
        user_id = session.get('user_id')

        print(f"\n{'='*50}")
        print(f"Upload request received. Session ID: {session_id}")
        print(f"Request files keys: {list(request.files.keys())}")
        print(f"Content-Length: {request.content_length}")

        # Check subscription limits
        if user_id:
            can_upload, current, limit = check_usage_limit(user_id, 'documents')
            if not can_upload:
                return jsonify({
                    'success': False,
                    'message': f'Document limit reached ({current}/{limit}). Upgrade to Pro for more uploads.',
                    'upgrade_required': True,
                    'current_usage': current,
                    'limit': limit
                })

            # Get user's plan for format restrictions
            plan_name = get_user_plan(user_id)
            plan = SUBSCRIPTION_PLANS[plan_name]
            allowed_formats = plan['allowed_formats']
        else:
            allowed_formats = ['pdf', 'docx', 'txt']

        if 'file' not in request.files and 'files' not in request.files:
            print("No file field found in request")
            return jsonify({'success': False, 'message': 'No file provided'})

        # Handle both single file and multiple files
        if 'files' in request.files:
            files = request.files.getlist('files')
        else:
            files = [request.files['file']]

        print(f"Files to process: {[f.filename for f in files]}")

        total_chunks = 0
        processed_files = []
        errors = []

        for file in files:
            if file.filename == '':
                continue

            try:
                filename = secure_filename(file.filename)
                ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

                print(f"\nProcessing: {filename} (type: {ext})")

                # Check if file format is allowed for user's plan
                if ext not in allowed_formats:
                    plan_name = get_user_plan(user_id) if user_id else 'free'
                    errors.append(f"{filename}: {ext.upper()} files not allowed on {plan_name.title()} plan. Upgrade to Pro.")
                    continue

                # Read file bytes
                file_bytes = file.read()
                file_size_mb = len(file_bytes) / (1024 * 1024)
                print(f"File size: {file_size_mb:.2f} MB")

                # Extract text based on file type
                text = ""
                detected_lang = "en"
                page_count = 0

                if ext == 'pdf':
                    print(f"Extracting text from PDF (Hindi + English OCR)...")
                    result = extract_text_from_pdf(file_bytes)
                    if isinstance(result, tuple):
                        text, detected_lang = result
                    else:
                        text = result

                    # Get page count for PDF
                    try:
                        pdf_fitz = get_fitz()
                        pdf_doc = pdf_fitz.open(stream=file_bytes, filetype="pdf")
                        page_count = len(pdf_doc)
                        pdf_doc.close()

                        # Check page limit for user's plan
                        if user_id:
                            plan_name = get_user_plan(user_id)
                            plan = SUBSCRIPTION_PLANS[plan_name]
                            max_pages = plan['max_pages']
                            if max_pages != -1 and page_count > max_pages:
                                errors.append(f"{filename}: PDF has {page_count} pages (max {max_pages} on {plan_name.title()} plan). Upgrade to Pro.")
                                continue
                    except:
                        pass

                elif ext == 'docx':
                    # Check if DOCX is allowed
                    if user_id and not check_feature_access(user_id, 'docx_support'):
                        errors.append(f"{filename}: DOCX files require Pro plan. Upgrade to access.")
                        continue

                    print(f"Extracting text from DOCX...")
                    result = extract_text_from_docx(file_bytes)
                    if isinstance(result, tuple):
                        text, detected_lang = result
                    else:
                        text = result
                elif ext == 'txt':
                    print(f"Reading TXT file...")
                    try:
                        text = file_bytes.decode('utf-8')
                    except:
                        text = file_bytes.decode('latin-1')
                    detected_lang = detect_language(text)
                else:
                    errors.append(f"{filename}: Unsupported file type")
                    continue

                # Check Hindi language support
                if detected_lang == 'hi' and user_id and not check_feature_access(user_id, 'hindi_support'):
                    errors.append(f"{filename}: Hindi documents require Pro plan. Upgrade to access.")
                    continue

                if not text or not text.strip():
                    error_msg = f"{filename}: Could not extract text (PDF might be scanned/image-based)"
                    print(error_msg)
                    errors.append(error_msg)
                    continue

                lang_name = "Hindi" if detected_lang == "hi" else "English"
                print(f"Extracted {len(text)} characters from {filename} (Language: {lang_name})")

                # Add to vector store with language info
                chunks_added = add_to_vector_store(session_id, text, filename, detected_lang)
                total_chunks += chunks_added
                processed_files.append(filename)

                # Track usage and add to history
                if user_id:
                    increment_usage(user_id, 'documents_uploaded', 1)
                    if page_count > 0:
                        increment_usage(user_id, 'pages_processed', page_count)

                    # Add to history if feature is available
                    if check_feature_access(user_id, 'history'):
                        add_to_history(user_id, {
                            'filename': filename,
                            'pages': page_count,
                            'language': detected_lang
                        })

                print(f"Successfully processed {filename}: {chunks_added} chunks")

            except Exception as e:
                error_msg = f"{file.filename}: {str(e)}"
                print(f"Error processing file: {error_msg}")
                traceback.print_exc()
                errors.append(error_msg)

        print(f"\n{'='*50}")
        print(f"Upload complete: {len(processed_files)} files, {total_chunks} total chunks")

        if processed_files:
            message = f"Successfully processed {len(processed_files)} file(s)"
            if errors:
                message += f". Errors: {'; '.join(errors)}"

            return jsonify({
                'success': True,
                'message': message,
                'chunks': total_chunks,
                'files': processed_files,
                'filename': processed_files[0] if len(processed_files) == 1 else None,
                'errors': errors
            })
        else:
            error_msg = '; '.join(errors) if errors else 'Unknown error - check if PDF is image-based'
            return jsonify({
                'success': False,
                'message': f"Failed to process files. {error_msg}"
            })

    except Exception as e:
        print(f"Error in upload_document: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'Upload error: {str(e)}'})

@app.route('/api/upload-url', methods=['POST'])
def upload_url():
    """Load content from URL(s)"""
    try:
        session_id = get_or_create_session_id()
        data = request.json
        urls = data.get('urls', [])

        if isinstance(urls, str):
            urls = [urls]

        if not urls:
            return jsonify({'success': False, 'message': 'No URLs provided'})

        total_chunks = 0
        processed_urls = []
        errors = []

        for url in urls:
            url = url.strip()
            if not url:
                continue

            try:
                # Validate URL
                parsed = urlparse(url)
                if not parsed.scheme:
                    url = 'https://' + url

                text = extract_text_from_url(url)

                if not text or not text.strip():
                    errors.append(f"{url}: Could not extract content")
                    continue

                # Add to vector store
                source_name = parsed.netloc or url[:50]
                chunks_added = add_to_vector_store(session_id, text, source_name)
                total_chunks += chunks_added
                processed_urls.append(url)

                print(f"Processed URL {url}: {chunks_added} chunks")

            except Exception as e:
                print(f"Error processing URL {url}: {e}")
                errors.append(f"{url}: {str(e)}")

        if processed_urls:
            return jsonify({
                'success': True,
                'message': f"Successfully loaded {len(processed_urls)} URL(s)",
                'chunks': total_chunks,
                'urls': processed_urls,
                'errors': errors
            })
        else:
            return jsonify({
                'success': False,
                'message': f"Failed to load URLs. {'; '.join(errors) if errors else ''}"
            })

    except Exception as e:
        print(f"Error in upload_url: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/chat', methods=['POST'])
def chat():
    try:
        data = request.json
        query = data.get('query', '').strip()
        output_language = data.get('output_language', 'en')  # 'en', 'hi', or 'both'
        api_key = session.get('api_key')
        session_id = session.get('session_id')
        user_id = session.get('user_id')

        if not api_key:
            return jsonify({'success': False, 'message': 'Please set your API key first'})

        if not query:
            return jsonify({'success': False, 'message': 'No query provided'})

        # Check semantic Q&A access for Free plan users
        if user_id and not check_feature_access(user_id, 'semantic_qa'):
            return jsonify({
                'success': False,
                'message': 'Semantic Q&A is a Pro feature. Upgrade to ask questions about your documents.',
                'upgrade_required': True
            })

        # Check Hindi output support
        if output_language in ['hi', 'both'] and user_id and not check_feature_access(user_id, 'hindi_support'):
            return jsonify({
                'success': False,
                'message': 'Hindi language support requires Pro plan. Upgrade to access.',
                'upgrade_required': True
            })

        # Get relevant context if documents are uploaded
        context = ""
        sources_used = []
        doc_language = 'en'

        if session_id and session_id in vector_stores:
            store = vector_stores[session_id]
            if store['chunks']:
                model = get_embedding_model()
                query_embedding = np.array(model.encode([query])).astype('float32')

                # Get more chunks for better context (up to 15)
                k = min(15, len(store['chunks']))
                distances, indices = store['index'].search(query_embedding, k)

                relevant_chunks = [store['chunks'][i] for i in indices[0] if i < len(store['chunks'])]
                context = "\n\n---\n\n".join(relevant_chunks)
                sources_used = list(set(store['sources']))

                # Get document language
                if 'languages' in store and store['languages']:
                    doc_language = list(store['languages'].values())[0]

                # Also include the beginning of the document for case details
                if store['texts']:
                    first_doc = list(store['texts'].values())[0]
                    # Add first 3000 chars which usually have case title, citation, etc.
                    doc_header = first_doc[:3000]
                    context = f"DOCUMENT HEADER:\n{doc_header}\n\n---\n\nRELEVANT SECTIONS:\n{context}"

        # Language instruction
        lang_instruction = ""
        if output_language == 'hi':
            lang_instruction = "\n\nIMPORTANT: Respond in Hindi (Devanagari script)."
        elif output_language == 'both':
            lang_instruction = "\n\nIMPORTANT: Provide your response in BOTH English and Hindi. First give the English response, then provide the same in Hindi (Devanagari script) under a '--- हिंदी में ---' separator."

        # Build prompt - focus on accurate, concise responses
        if context:
            prompt = f"""You are JurisMind AI, an expert legal assistant. Answer the user's question based ONLY on the document context provided.

STRICT RULES:
1. Give DIRECT, ACCURATE answers - no roundabout explanations
2. For factual questions (names, dates, sections, citations), give the EXACT answer in 1-2 sentences
3. Only elaborate if the question requires explanation
4. If the answer is not in the document, say "This information is not found in the document"
5. NEVER make up information - only use what's in the document

DOCUMENT CONTEXT:
{context}

USER QUESTION: {query}

Provide a precise, accurate answer. For simple factual questions, give a direct 1-2 line response. For complex questions, structure your answer clearly but concisely.{lang_instruction}"""
        else:
            prompt = f"""You are JurisMind AI, an expert legal assistant. Answer the following legal question.

STRICT RULES:
1. Give DIRECT, ACCURATE answers - no roundabout explanations
2. Be concise - 1-2 sentences for simple questions
3. Note that this is general legal information, not advice

USER QUESTION: {query}

Provide a precise, helpful answer.{lang_instruction}"""

        response = gemini_response(prompt, api_key)

        if response:
            # Increment Q&A counter
            if user_id:
                increment_usage(user_id, 'qa_queries')

            # Save Q&A to history if user has history access and there's context
            if user_id and context and check_feature_access(user_id, 'history'):
                doc_id = get_current_doc_id(user_id)
                if doc_id:
                    add_chat_to_history(user_id, doc_id, query, response)

            return jsonify({
                'success': True,
                'response': response,
                'sources': sources_used,
                'document_language': doc_language
            })
        else:
            return jsonify({'success': False, 'message': 'Failed to get AI response. Please check your API key.'})

    except Exception as e:
        print(f"Error in chat: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/generate-brief', methods=['POST'])
def generate_brief():
    try:
        data = request.json or {}
        output_language = data.get('output_language', 'en')  # 'en', 'hi', or 'both'
        api_key = session.get('api_key')
        session_id = session.get('session_id')

        if not api_key:
            return jsonify({'success': False, 'message': 'Please set your API key first'})

        if not session_id or session_id not in vector_stores:
            return jsonify({'success': False, 'message': 'Please upload a document first'})

        store = vector_stores[session_id]

        # Get document language
        doc_language = 'en'
        if 'languages' in store and store['languages']:
            doc_language = list(store['languages'].values())[0]

        # Get all text from uploaded documents
        all_text = "\n\n".join(store['texts'].values())
        # Send more text for better case brief (up to 50000 chars for comprehensive extraction)
        text = all_text[:50000]

        # Language instruction
        lang_instruction = ""
        if output_language == 'hi':
            lang_instruction = "\n\nOUTPUT LANGUAGE: Generate the entire case brief in Hindi (Devanagari script). Use Hindi legal terminology where appropriate."
        elif output_language == 'both':
            lang_instruction = "\n\nOUTPUT LANGUAGE: Generate the case brief in BOTH English and Hindi. First provide the complete English version, then after a '═══════════════════════════════════════\n📋 हिंदी में केस ब्रीफ (Case Brief in Hindi)\n═══════════════════════════════════════' separator, provide the complete Hindi version in Devanagari script."

        prompt = f"""You are JurisMind AI, an expert legal assistant specialized in Indian law. Generate a comprehensive and detailed case brief for the following legal document.

CRITICAL INSTRUCTIONS:
1. Extract EVERY specific detail mentioned in the document - do NOT generalize or summarize loosely
2. Include ALL exact names, dates, numbers, section references as they appear in the document
3. If information for any section is not found, explicitly state "Not mentioned in the document"
4. Quote key phrases directly from the document where relevant

DOCUMENT TEXT:
{text}

Generate a professional case brief with the following sections:

## 📋 CASE IDENTIFICATION
- **Case Title**: [EXACT case name as mentioned, e.g., "Satender Kumar Antil vs. Central Bureau of Investigation"]
- **Citation**: [Full citation, e.g., "(2022) 10 SCC 51" or "2022 SCC OnLine SC 825"]
- **Court**: [Court name, e.g., "Supreme Court of India"]
- **Case Number**: [Criminal Appeal No./Writ Petition No./SLP No. with year]
- **Date of Judgment**: [Exact date]

## 👨‍⚖️ BENCH (CORAM)
- List ALL judges by name who heard the case
- Indicate who authored the judgment if mentioned

## 👥 PARTIES
- **Appellant/Petitioner**: [Full name(s) with any designation]
- **Respondent**: [Full name(s) with any designation]
- **Advocates**: [Names of lawyers appearing for each side if mentioned]

## 📜 FACTS OF THE CASE
- Present the factual background in chronological order
- Include specific dates, places, amounts, and circumstances
- Mention the original complaint/FIR details if any

## ⚖️ PROCEDURAL HISTORY
- Track the case through different courts/forums
- Include dates and outcomes at each stage
- Mention the specific order being challenged

## ❓ ISSUES/QUESTIONS OF LAW
- List each legal question the court addressed
- Include relevant statutory provisions (e.g., Section 438 CrPC, Section 302 IPC)
- Number each issue clearly

## 🏛️ ARGUMENTS
**Appellant's Arguments:**
- Key contentions raised

**Respondent's Arguments:**
- Key contentions in response

## 📝 HOLDING/DECISION
- State the court's decision on EACH issue
- Include the final order/directions given
- Mention any relief granted or denied

## 💡 REASONING/RATIO DECIDENDI
- Explain the legal reasoning behind each decision
- Include key observations and interpretations
- Quote significant passages from the judgment

## 📚 LEGAL PROVISIONS DISCUSSED
- List ALL sections, articles, and rules mentioned (e.g., Section 170, 438, 468 CrPC, Article 21)
- Briefly note how each was applied

## 🔖 PRECEDENTS CITED
- List cases referred to by the court
- Note how they were applied or distinguished

## ⭐ SIGNIFICANCE/IMPLICATIONS
- Legal principles established
- Impact on future cases
- Practical implications for legal practice{lang_instruction}

Format the brief professionally with clear headings and markdown formatting. Be thorough and accurate."""

        response = gemini_response(prompt, api_key)

        if response:
            # Save brief to history if user has history access
            user_id = session.get('user_id')
            if user_id:
                # Increment briefs_generated counter
                increment_usage(user_id, 'briefs_generated')

                # Save to history if user has access
                if check_feature_access(user_id, 'history'):
                    doc_id = get_current_doc_id(user_id)
                    if doc_id:
                        update_history_brief(user_id, doc_id, response)

            return jsonify({
                'success': True,
                'brief': response,
                'document_language': doc_language
            })
        else:
            return jsonify({'success': False, 'message': 'Failed to generate brief'})

    except Exception as e:
        print(f"Error in generate_brief: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/analyze', methods=['POST'])
def analyze_document():
    try:
        api_key = session.get('api_key')
        session_id = session.get('session_id')
        user_id = session.get('user_id')

        if not api_key:
            return jsonify({'success': False, 'message': 'Please set your API key first'})

        if not session_id or session_id not in vector_stores:
            return jsonify({'success': False, 'message': 'Please upload a document first'})

        # Check document analysis access
        if user_id and not check_feature_access(user_id, 'document_analysis'):
            return jsonify({
                'success': False,
                'message': 'Document Analysis is a Pro feature. Upgrade to access detailed document insights.',
                'upgrade_required': True
            })

        store = vector_stores[session_id]

        # Get all text from uploaded documents
        all_text = "\n\n".join(store['texts'].values())
        text = all_text[:20000]

        prompt = f"""You are JurisMind AI, an expert legal analyst. Analyze the following legal document and provide:

DOCUMENT:
{text}

Provide a comprehensive analysis including:
1. **Document Type** - What kind of legal document is this?
2. **Key Parties** - Who are the parties involved?
3. **Main Legal Issues** - What legal matters are addressed?
4. **Important Clauses/Sections** - Highlight critical parts
5. **Potential Risks** - Any concerns or red flags
6. **Recommendations** - Suggested actions or considerations
7. **Summary** - Brief overview of the document

Be thorough but concise. Use markdown formatting for better readability."""

        response = gemini_response(prompt, api_key)

        if response:
            return jsonify({'success': True, 'analysis': response})
        else:
            return jsonify({'success': False, 'message': 'Failed to analyze document'})

    except Exception as e:
        print(f"Error in analyze_document: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/clear', methods=['POST'])
def clear_documents():
    """Clear all uploaded documents for the session"""
    try:
        session_id = session.get('session_id')
        if session_id and session_id in vector_stores:
            del vector_stores[session_id]
        return jsonify({'success': True, 'message': 'All documents cleared'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/status', methods=['GET'])
def get_status():
    """Get current session status"""
    try:
        session_id = session.get('session_id')

        # Auto-set API key from .env if not already set
        if 'api_key' not in session:
            env_api_key = os.getenv('GEMINI_API_KEY')
            if env_api_key:
                session['api_key'] = env_api_key

        api_key_set = bool(session.get('api_key'))

        docs_count = 0
        chunks_count = 0
        sources = []

        if session_id and session_id in vector_stores:
            store = vector_stores[session_id]
            docs_count = len(store.get('texts', {}))
            chunks_count = len(store.get('chunks', []))
            sources = store.get('sources', [])

        return jsonify({
            'success': True,
            'api_key_set': api_key_set,
            'documents_count': docs_count,
            'chunks_count': chunks_count,
            'sources': sources
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/check-feature-access', methods=['GET'])
def check_feature_access_api():
    """Check if current user has access to a specific feature"""
    try:
        user_id = session.get('user_id')
        feature = request.args.get('feature')

        if not feature:
            return jsonify({'success': False, 'message': 'Feature not specified'})

        if not user_id:
            return jsonify({'success': True, 'has_access': False})

        has_access = check_feature_access(user_id, feature)
        return jsonify({'success': True, 'has_access': has_access})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e), 'has_access': False})

# ==================== SUBSCRIPTION API ENDPOINTS ====================

@app.route('/api/subscription/plans', methods=['GET'])
def get_plans():
    """Get all subscription plans"""
    return jsonify({
        'success': True,
        'plans': SUBSCRIPTION_PLANS
    })

@app.route('/api/subscription/current', methods=['GET'])
def get_current_subscription():
    """Get current user's subscription details"""
    try:
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'success': False, 'message': 'Not logged in'})

        plan_name = get_user_plan(user_id)
        plan = SUBSCRIPTION_PLANS[plan_name]
        usage = get_user_usage(user_id)

        # Calculate remaining limits
        docs_limit = plan['documents_per_month']
        pages_limit = plan['max_pages']
        docs_used = usage.get('documents_uploaded', 0)
        pages_used = usage.get('pages_processed', 0)

        return jsonify({
            'success': True,
            'plan': plan_name,
            'plan_details': plan,
            'usage': usage,
            'limits': {
                'documents': {
                    'used': docs_used,
                    'limit': docs_limit,
                    'remaining': docs_limit - docs_used if docs_limit != -1 else -1
                },
                'pages': {
                    'used': pages_used,
                    'limit': pages_limit,
                    'remaining': pages_limit - pages_used if pages_limit != -1 else -1
                }
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/subscription/upi-info', methods=['GET'])
def get_upi_info():
    """Get UPI payment information"""
    return jsonify({
        'success': True,
        'upi_id': UPI_ID,
        'amount': SUBSCRIPTION_PLANS['pro']['price_inr'],
        'plan_name': 'Pro'
    })

@app.route('/api/subscription/request-upgrade', methods=['POST'])
def request_upgrade():
    """Submit payment request with screenshot"""
    try:
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'success': False, 'message': 'Not logged in'})

        # Check if screenshot is provided
        if 'screenshot' not in request.files:
            return jsonify({'success': False, 'message': 'Please upload payment screenshot'})

        screenshot = request.files['screenshot']
        transaction_id = request.form.get('transaction_id', '').strip()

        if not transaction_id:
            return jsonify({'success': False, 'message': 'Please enter transaction ID'})

        if screenshot.filename == '':
            return jsonify({'success': False, 'message': 'Please select a screenshot file'})

        # Validate file type
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif'}
        ext = screenshot.filename.rsplit('.', 1)[-1].lower() if '.' in screenshot.filename else ''
        if ext not in allowed_extensions:
            return jsonify({'success': False, 'message': 'Invalid file type. Please upload an image (PNG, JPG, JPEG, GIF)'})

        # Save screenshot
        filename = f"{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"
        filepath = os.path.join(PAYMENT_SCREENSHOTS_DIR, filename)
        screenshot.save(filepath)

        # Load existing requests
        payment_requests = load_payment_requests()

        # Get user info
        users = load_users_db()
        user_info = users.get(user_id, {})

        # Create payment request
        request_id = secrets.token_hex(8)
        payment_requests[request_id] = {
            'user_id': user_id,
            'user_name': user_info.get('name', 'Unknown'),
            'user_email': user_info.get('email', 'Unknown'),
            'user_phone': user_info.get('phone', 'Unknown'),
            'transaction_id': transaction_id,
            'screenshot_path': filename,
            'amount': SUBSCRIPTION_PLANS['pro']['price_inr'],
            'plan': 'pro',
            'status': 'pending',  # pending, approved, rejected
            'submitted_at': datetime.now().isoformat(),
            'reviewed_at': None,
            'reviewed_by': None
        }

        save_payment_requests(payment_requests)

        return jsonify({
            'success': True,
            'message': 'Payment request submitted successfully! Your request will be reviewed within 24 hours.',
            'request_id': request_id
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/subscription/request-status', methods=['GET'])
def get_request_status():
    """Get user's payment request status"""
    try:
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'success': False, 'message': 'Not logged in'})

        payment_requests = load_payment_requests()

        # Find user's requests
        user_requests = []
        for req_id, req in payment_requests.items():
            if req['user_id'] == user_id:
                user_requests.append({
                    'request_id': req_id,
                    'status': req['status'],
                    'submitted_at': req['submitted_at'],
                    'reviewed_at': req.get('reviewed_at'),
                    'transaction_id': req['transaction_id']
                })

        # Sort by submission date (newest first)
        user_requests.sort(key=lambda x: x['submitted_at'], reverse=True)

        return jsonify({
            'success': True,
            'requests': user_requests,
            'has_pending': any(r['status'] == 'pending' for r in user_requests)
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

# ==================== ADMIN ENDPOINTS ====================

@app.route('/admin')
def admin_page():
    """Admin dashboard page"""
    return render_template('admin.html')

@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    """Admin login"""
    try:
        data = request.json
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')

        if email == ADMIN_EMAIL.lower() and password == ADMIN_PASSWORD:
            session['is_admin'] = True
            session['admin_email'] = ADMIN_EMAIL
            return jsonify({'success': True, 'message': 'Admin login successful'})
        else:
            return jsonify({'success': False, 'message': 'Invalid email or password'})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/pending-requests', methods=['GET'])
def get_pending_requests():
    """Get all pending payment requests (admin only)"""
    try:
        if not session.get('is_admin'):
            return jsonify({'success': False, 'message': 'Admin access required'})

        payment_requests = load_payment_requests()

        # Get all requests with details
        all_requests = []
        for req_id, req in payment_requests.items():
            all_requests.append({
                'request_id': req_id,
                **req
            })

        # Sort by status (pending first) and date
        all_requests.sort(key=lambda x: (0 if x['status'] == 'pending' else 1, x['submitted_at']), reverse=True)

        return jsonify({
            'success': True,
            'requests': all_requests,
            'pending_count': sum(1 for r in all_requests if r['status'] == 'pending')
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/approve-request', methods=['POST'])
def approve_request():
    """Approve a payment request (admin only)"""
    try:
        if not session.get('is_admin'):
            return jsonify({'success': False, 'message': 'Admin access required'})

        data = request.json
        request_id = data.get('request_id')

        if not request_id:
            return jsonify({'success': False, 'message': 'Request ID required'})

        payment_requests = load_payment_requests()

        if request_id not in payment_requests:
            return jsonify({'success': False, 'message': 'Request not found'})

        req = payment_requests[request_id]

        if req['status'] != 'pending':
            return jsonify({'success': False, 'message': 'Request already processed'})

        # Approve the request
        payment_requests[request_id]['status'] = 'approved'
        payment_requests[request_id]['reviewed_at'] = datetime.now().isoformat()
        payment_requests[request_id]['reviewed_by'] = 'admin'
        save_payment_requests(payment_requests)

        # Upgrade user's plan
        users = load_users_db()
        user_id = req['user_id']
        if user_id in users:
            users[user_id]['plan'] = 'pro'
            users[user_id]['plan_updated'] = datetime.now().isoformat()
            users[user_id]['plan_expires'] = (datetime.now() + timedelta(days=30)).isoformat()
            save_users_db(users)

        return jsonify({
            'success': True,
            'message': f'Payment approved! User {req["user_name"]} upgraded to Pro plan.'
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/reject-request', methods=['POST'])
def reject_request():
    """Reject a payment request (admin only)"""
    try:
        if not session.get('is_admin'):
            return jsonify({'success': False, 'message': 'Admin access required'})

        data = request.json
        request_id = data.get('request_id')
        reason = data.get('reason', 'Payment verification failed')

        if not request_id:
            return jsonify({'success': False, 'message': 'Request ID required'})

        payment_requests = load_payment_requests()

        if request_id not in payment_requests:
            return jsonify({'success': False, 'message': 'Request not found'})

        if payment_requests[request_id]['status'] != 'pending':
            return jsonify({'success': False, 'message': 'Request already processed'})

        # Reject the request
        payment_requests[request_id]['status'] = 'rejected'
        payment_requests[request_id]['reviewed_at'] = datetime.now().isoformat()
        payment_requests[request_id]['reviewed_by'] = 'admin'
        payment_requests[request_id]['rejection_reason'] = reason
        save_payment_requests(payment_requests)

        return jsonify({
            'success': True,
            'message': 'Payment request rejected.'
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/screenshot/<filename>')
def get_screenshot(filename):
    """Serve payment screenshot (admin only)"""
    if not session.get('is_admin'):
        return "Unauthorized", 403

    from flask import send_from_directory
    return send_from_directory(PAYMENT_SCREENSHOTS_DIR, filename)

@app.route('/api/subscription/history', methods=['GET'])
def get_document_history():
    """Get user's document history with full details (briefs and Q&A)"""
    try:
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'success': False, 'message': 'Not logged in'})

        # Check if user has access to history feature
        if not check_feature_access(user_id, 'history'):
            return jsonify({
                'success': False,
                'message': 'Document history is a Pro feature. Upgrade to access.',
                'upgrade_required': True
            })

        limit = request.args.get('limit', 20, type=int)
        include_details = request.args.get('details', 'true').lower() == 'true'
        history = get_user_history(user_id, limit)

        # Format history for frontend
        formatted_history = []
        for item in history:
            filename = item.get('filename', 'Document')
            file_type = 'pdf'
            if filename:
                ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
                if ext in ['docx', 'doc']:
                    file_type = 'docx'
                elif ext == 'txt':
                    file_type = 'txt'

            history_entry = {
                'id': item.get('id'),
                'filename': filename,
                'uploaded_at': item.get('upload_date'),
                'pages': item.get('pages', 0),
                'type': file_type,
                'language': item.get('language', 'en'),
                'brief_generated': item.get('brief_generated', False),
                'analysis_done': item.get('analysis_done', False)
            }

            # Include full details if requested
            if include_details:
                history_entry['case_brief'] = item.get('case_brief')
                history_entry['brief_generated_at'] = item.get('brief_generated_at')
                history_entry['chat_history'] = item.get('chat_history', [])

            formatted_history.append(history_entry)

        return jsonify({
            'success': True,
            'history': formatted_history
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/export/brief', methods=['POST'])
def export_brief():
    """Export case brief as PDF or DOCX"""
    try:
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'success': False, 'message': 'Not logged in'})

        data = request.json
        export_format = data.get('format', 'pdf').lower()
        brief_content = data.get('content', '')

        if not brief_content:
            return jsonify({'success': False, 'message': 'No content to export'})

        # Check feature access
        feature = 'export_pdf' if export_format == 'pdf' else 'export_docx'
        if not check_feature_access(user_id, feature):
            return jsonify({
                'success': False,
                'message': f'Export to {export_format.upper()} is a Pro feature. Upgrade to access.',
                'upgrade_required': True
            })

        if export_format == 'pdf':
            # Generate PDF
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            import tempfile
            import base64

            # Create temp file path
            temp_fd, temp_path = tempfile.mkstemp(suffix='.pdf')
            os.close(temp_fd)  # Close the file descriptor

            try:
                doc = SimpleDocTemplate(temp_path, pagesize=A4,
                                      rightMargin=72, leftMargin=72,
                                      topMargin=72, bottomMargin=18)

                styles = getSampleStyleSheet()
                story = []

                # Add title
                title_style = ParagraphStyle('Title', parent=styles['Heading1'],
                                            fontSize=18, spaceAfter=30)
                story.append(Paragraph("Case Brief - JurisMind AI", title_style))
                story.append(Spacer(1, 0.25*inch))

                # Add content (convert markdown-like content to paragraphs)
                content_lines = brief_content.split('\n')
                for line in content_lines:
                    if line.strip():
                        # Escape HTML special characters
                        safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        if line.startswith('## '):
                            story.append(Spacer(1, 0.2*inch))
                            story.append(Paragraph(safe_line[3:], styles['Heading2']))
                        elif line.startswith('### '):
                            story.append(Spacer(1, 0.15*inch))
                            story.append(Paragraph(safe_line[4:], styles['Heading3']))
                        elif line.startswith('**') and line.endswith('**'):
                            story.append(Paragraph(f"<b>{safe_line[2:-2]}</b>", styles['Normal']))
                        elif line.startswith('- '):
                            story.append(Paragraph(f"• {safe_line[2:]}", styles['Normal']))
                        else:
                            story.append(Paragraph(safe_line, styles['Normal']))
                        story.append(Spacer(1, 0.05*inch))

                doc.build(story)

                # Read file and return
                with open(temp_path, 'rb') as f:
                    pdf_content = f.read()

                pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')
            finally:
                # Clean up temp file
                try:
                    os.unlink(temp_path)
                except:
                    pass

            return jsonify({
                'success': True,
                'format': 'pdf',
                'filename': 'case_brief.pdf',
                'content': pdf_base64
            })

        elif export_format == 'docx':
            # Generate DOCX
            from docx import Document
            from docx.shared import Inches, Pt
            import tempfile
            import base64

            # Create temp file path
            temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
            os.close(temp_fd)

            try:
                doc = Document()
                doc.add_heading('Case Brief - JurisMind AI', 0)

                # Add content
                content_lines = brief_content.split('\n')
                for line in content_lines:
                    if line.strip():
                        if line.startswith('## '):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith('### '):
                            doc.add_heading(line[4:], level=3)
                        elif line.startswith('- '):
                            doc.add_paragraph(line[2:], style='List Bullet')
                        elif line.startswith('**') and line.endswith('**'):
                            p = doc.add_paragraph()
                            p.add_run(line[2:-2]).bold = True
                        else:
                            doc.add_paragraph(line)

                doc.save(temp_path)

                with open(temp_path, 'rb') as f:
                    docx_content = f.read()

                docx_base64 = base64.b64encode(docx_content).decode('utf-8')
            finally:
                try:
                    os.unlink(temp_path)
                except:
                    pass

            return jsonify({
                'success': True,
                'format': 'docx',
                'filename': 'case_brief.docx',
                'content': docx_base64
            })

        else:
            return jsonify({'success': False, 'message': 'Invalid export format'})

    except ImportError as e:
        return jsonify({'success': False, 'message': f'Export library not installed: {str(e)}'})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

# ==================== GAMIFICATION ENDPOINTS ====================

@app.route('/api/gamification/stats', methods=['GET'])
def get_gamification_stats():
    """Get user's complete gamification stats including usage, achievements, and daily tip"""
    try:
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'success': False, 'message': 'Not logged in'})

        # Get total stats
        total_stats = get_total_user_stats(user_id)

        # Get current month stats
        current_stats = get_user_usage(user_id)

        # Get achievements
        user_achievements = get_user_achievements(user_id)
        unlocked_ids = user_achievements.get('unlocked', [])

        # Build achievement list with status
        all_achievements = []
        for ach_id, achievement in ACHIEVEMENTS.items():
            ach_data = {
                **achievement,
                'unlocked': ach_id in unlocked_ids,
                'progress': min(100, int((total_stats.get(achievement['field'], 0) / achievement['requirement']) * 100))
            }
            all_achievements.append(ach_data)

        # Sort achievements: unlocked first, then by progress
        all_achievements.sort(key=lambda x: (-x['unlocked'], -x['progress']))

        # Check for new achievements
        new_achievements = check_and_unlock_achievements(user_id, total_stats)

        # Get daily tip
        daily_tip = get_daily_tip()

        # Calculate time saved (estimate: 30 min per brief, 5 min per question)
        time_saved_minutes = (total_stats['briefs_generated'] * 30) + (total_stats['qa_queries'] * 5)
        time_saved_hours = time_saved_minutes / 60

        # Get user's join date for "member since"
        users = load_users_db()
        member_since = None
        if user_id in users:
            created_at = users[user_id].get('created_at')
            if created_at:
                member_since = datetime.fromisoformat(created_at).strftime('%B %Y')

        return jsonify({
            'success': True,
            'total_stats': total_stats,
            'current_month_stats': current_stats,
            'achievements': all_achievements,
            'new_achievements': new_achievements,
            'daily_tip': daily_tip,
            'time_saved': {
                'minutes': time_saved_minutes,
                'hours': round(time_saved_hours, 1),
                'display': f"{int(time_saved_hours)}h {int(time_saved_minutes % 60)}m" if time_saved_hours >= 1 else f"{time_saved_minutes}m"
            },
            'member_since': member_since,
            'unlocked_count': len(unlocked_ids),
            'total_achievements': len(ACHIEVEMENTS)
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/gamification/achievements', methods=['GET'])
def get_achievements():
    """Get all achievements with user's progress"""
    try:
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'success': False, 'message': 'Not logged in'})

        total_stats = get_total_user_stats(user_id)
        user_achievements = get_user_achievements(user_id)
        unlocked_ids = user_achievements.get('unlocked', [])

        achievements_list = []
        for ach_id, achievement in ACHIEVEMENTS.items():
            current_value = total_stats.get(achievement['field'], 0)
            achievements_list.append({
                **achievement,
                'unlocked': ach_id in unlocked_ids,
                'current_value': current_value,
                'progress': min(100, int((current_value / achievement['requirement']) * 100))
            })

        # Sort: unlocked first, then by progress descending
        achievements_list.sort(key=lambda x: (-x['unlocked'], -x['progress']))

        return jsonify({
            'success': True,
            'achievements': achievements_list,
            'unlocked_count': len(unlocked_ids),
            'total_count': len(ACHIEVEMENTS)
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/gamification/daily-tip', methods=['GET'])
def get_tip():
    """Get today's legal tip"""
    try:
        tip = get_daily_tip()
        return jsonify({
            'success': True,
            'tip': tip
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/gamification/mark-achievement-seen', methods=['POST'])
def mark_achievement_seen():
    """Mark new achievements as seen/notified"""
    try:
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'success': False, 'message': 'Not logged in'})

        data = request.json
        achievement_ids = data.get('achievement_ids', [])

        achievements_db = load_achievements_db()
        if user_id not in achievements_db:
            achievements_db[user_id] = {'unlocked': [], 'notified': []}

        notified = achievements_db[user_id].get('notified', [])
        for ach_id in achievement_ids:
            if ach_id not in notified:
                notified.append(ach_id)

        achievements_db[user_id]['notified'] = notified
        save_achievements_db(achievements_db)

        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

# Pricing page route
@app.route('/pricing')
def pricing_page():
    """Pricing page"""
    return render_template('pricing.html')

if __name__ == '__main__':
    print("Starting JurisMind AI server...")
    print("Server ready! Embedding model will be loaded on first document upload.")
    # Disable reloader to prevent interruptions during uploads
    # Use PORT environment variable for Render, default to 5000 for local
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, port=port, threaded=True, use_reloader=False, host='0.0.0.0')
